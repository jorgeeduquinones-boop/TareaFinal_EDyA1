"""
Genera el documento Word de entrega para la Tarea Final EDyA1.
Sigue la estructura exigida en el enunciado:
  - Página 1: portada (universidad, facultad, asignatura, profesor, grupo,
    códigos y nombres de integrantes, fecha de entrega).
  - Páginas 2-5: resumen del enunciado (<=10 líneas), pseudocódigo de
    calcularPedidos con las complejidades individuales en columna lateral,
    y bloque aparte con la complejidad total.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "Tarea_Final_EDyA1_Entregable.docx"

doc = Document()

# Márgenes ajustados para entrar en 5 páginas sin apretar.
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# Estilo base
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def add_para(text, *, bold=False, italic=False, size=11, align=None, space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_heading(text, *, size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=10, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p

# =============================================================================
# PÁGINA 1 — PORTADA
# =============================================================================
add_para("UNIVERSIDAD AUTÓNOMA DE OCCIDENTE", bold=True, size=14,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para("Facultad de Ingeniería y Ciencias Básicas", size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para("Estructuras de Datos y Algoritmos 1 · 2026-1", size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

add_para("TAREA FINAL", bold=True, size=18,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("Análisis de Pedidos en Comercio Electrónico", bold=True, italic=True,
         size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

# Tabla con datos del grupo
tab = doc.add_table(rows=6, cols=2)
tab.style = "Light Grid Accent 1"
tab.autofit = False
tab.columns[0].width = Cm(5.5)
tab.columns[1].width = Cm(10.5)

filas = [
    ("Asignatura",  "Estructuras de Datos y Algoritmos 1"),
    ("Profesor",    "Orlando Arboleda Molina, Msc."),
    ("Facultad",    "Ingeniería y Ciencias Básicas"),
    ("Grupo EDyA1", "1"),
    ("Fecha de entrega", "17 de mayo de 2026 — 11:00 PM (UAOVirtual)"),
    ("Sustentación",     "Semana 16"),
]
for i, (k, v) in enumerate(filas):
    c0, c1 = tab.rows[i].cells
    c0.width = Cm(5.5)
    c1.width = Cm(10.5)
    c0.paragraphs[0].add_run(k).bold = True
    c1.paragraphs[0].add_run(v)
    for c in (c0, c1):
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

add_para("", space_after=10)

add_heading("Integrantes del grupo", size=13, space_before=10, space_after=4)

team = doc.add_table(rows=4, cols=2)
team.style = "Light Grid Accent 1"
team.autofit = False
team.columns[0].width = Cm(11.0)
team.columns[1].width = Cm(5.0)
encab = team.rows[0]
encab.cells[0].paragraphs[0].add_run("Nombre").bold = True
encab.cells[1].paragraphs[0].add_run("Código UAO").bold = True

integrantes = [
    ("Joan Mateo Cardona",     "2243431"),
    ("Danna Villegas",         "2240027"),
    ("Jorge Eduardo Álvarez",  "2230610"),
]
for i, (n, c) in enumerate(integrantes, start=1):
    team.rows[i].cells[0].paragraphs[0].add_run(n)
    team.rows[i].cells[1].paragraphs[0].add_run(c)
    for cell in team.rows[i].cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

add_para("", space_after=20)

p_repo = doc.add_paragraph()
p_repo.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_repo.add_run("Repositorio: ").bold = True
p_repo.add_run("github.com/jorgeeduquinones-boop/TareaFinal_EDyA1").italic = True

p_demo = doc.add_paragraph()
p_demo.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_demo.add_run("Demo en vivo: ").bold = True
p_demo.add_run("https://tarea-final-e-dy-a1.vercel.app/app/index.html").italic = True

# Salto de página → contenido en página 2
doc.add_page_break()

# =============================================================================
# PÁGINA 2 — RESUMEN DEL ENUNCIADO
# =============================================================================
add_heading("1. Resumen del enunciado", size=13, space_before=0)

resumen = (
    "Dado un caso compuesto por m records de compras (m ≥ 5) separados por "
    "punto y coma, donde cada record tiene el formato "
    "“customer product category price quantity timestamp”, se debe implementar "
    "la función calcularPedidos(caso) en util.js que retorne una cadena con n "
    "líneas (n ≥ 3, una por cliente distinto) en el formato "
    "“[a]) customer [totalSpent] [favoriteCategory]”. El ranking se ordena por "
    "tres criterios: (1) totalSpent — suma de price·quantity — descendente, "
    "(2) favoriteCategory — categoría con más transacciones por cliente — "
    "descendente, y (3) customer ascendente como desempate final. La solución "
    "debe usar algoritmos vistos en clase (Prácticas 4, 5 y 6) y procesar "
    "eficientemente entradas con alto volumen, conservando los términos n, m y p."
)
add_para(resumen, size=11, space_after=10)

# =============================================================================
# PSEUDOCÓDIGO + COMPLEJIDADES (tabla a dos columnas)
# =============================================================================
add_heading("2. Pseudocódigo de calcularPedidos(caso) y complejidades individuales",
            size=13, space_before=6, space_after=4)

intro = (
    "El pseudocódigo se presenta a la izquierda y las complejidades de cada "
    "fragmento al lado derecho, donde n = #clientes únicos, m = #records de "
    "entrada, p = #categorías distintas y k = longitud media de un campo."
)
add_para(intro, size=11, italic=True, space_after=6)

code_table = doc.add_table(rows=1, cols=2)
code_table.style = "Light Grid Accent 1"
code_table.autofit = False
code_table.columns[0].width = Cm(12.5)
code_table.columns[1].width = Cm(3.8)

hdr = code_table.rows[0].cells
hdr[0].paragraphs[0].add_run("Pseudocódigo").bold = True
hdr[1].paragraphs[0].add_run("Complejidad").bold = True

filas_codigo = [
    ("función calcularPedidos(caso):", ""),
    ("    si caso no es string o caso = '' → retornar ''", "O(1)"),
    ("    pedidos ← parsearCaso(caso)", "O(m · k)"),
    ("    si pedidos = ∅ → retornar ''", "O(1)"),
    ("    indice ← hash vacío;  clientes ← lista vacía", "O(1)"),
    ("    para cada ped en pedidos:                     // m iteraciones", "O(m)"),
    ("        si indice[ped.customer] = ∅:", ""),
    ("            cli ← nuevo Cliente(ped.customer)", ""),
    ("            indice[ped.customer] ← cli;  clientes.push(cli)", ""),
    ("        cli.registrarPedido(ped)                   // O(1) amortizado", ""),
    ("    para cada cli en clientes:                     // n iteraciones", "O(n · p)"),
    ("        cli.calcularCategoriaFavorita()            // O(p)", ""),
    ("    quickSort(clientes, 0, n−1, comparadorClientes)", "O(n log n) prom."),
    ("    salida ← ''", "O(1)"),
    ("    para i = 0 hasta n−1:", "O(n · k)"),
    ("        salida ← salida + (i+1) + ') ' + clientes[i].toString()", ""),
    ("    retornar salida", "O(1)"),
]

for code, comp in filas_codigo:
    row = code_table.add_row()
    c_code, c_comp = row.cells
    p_code = c_code.paragraphs[0]
    r_code = p_code.add_run(code)
    r_code.font.name = "Consolas"
    r_code.font.size = Pt(9.5)
    p_code.paragraph_format.space_after = Pt(0)

    p_comp = c_comp.paragraphs[0]
    r_comp = p_comp.add_run(comp)
    r_comp.font.name = "Consolas"
    r_comp.font.size = Pt(9.5)
    r_comp.bold = bool(comp)
    p_comp.paragraph_format.space_after = Pt(0)

# Página 3-4 continuación: complejidad total + observaciones

add_heading("3. Complejidad total", size=13, space_before=10, space_after=4)

# Tabla resumen aparte
tot = doc.add_table(rows=7, cols=3)
tot.style = "Light Grid Accent 1"
tot.autofit = False
tot.columns[0].width = Cm(5.5)
tot.columns[1].width = Cm(3.5)
tot.columns[2].width = Cm(7.3)

cab = tot.rows[0].cells
cab[0].paragraphs[0].add_run("Fragmento").bold = True
cab[1].paragraphs[0].add_run("Complejidad").bold = True
cab[2].paragraphs[0].add_run("Justificación").bold = True

fragmentos = [
    ("Parseo del caso",           "O(m · k)",   "Un solo recorrido del string de longitud L = m·k."),
    ("Acumulación por cliente",   "O(m)",       "Hash O(1) amortizado por record."),
    ("Categoría favorita",        "O(n · p)",   "Cada cliente recorre ≤ p categorías distintas."),
    ("QuickSort recursivo",       "O(n log n) promedio", "Mediana de tres + insertion sort umbral 16. Peor caso O(n²) acotado en la práctica."),
    ("Formateo de salida",        "O(n · k)",   "n strings de longitud O(k) concatenados."),
    ("TOTAL",                     "O(m·k + n·p + n log n)", "Suma de las etapas. Bajo k constante ≈ O(m + n·p + n log n)."),
]
for i, (frag, comp, just) in enumerate(fragmentos, start=1):
    cells = tot.rows[i].cells
    r0 = cells[0].paragraphs[0].add_run(frag)
    r1 = cells[1].paragraphs[0].add_run(comp)
    r2 = cells[2].paragraphs[0].add_run(just)
    if frag == "TOTAL":
        r0.bold = True; r1.bold = True

# Complejidad espacial
add_heading("4. Complejidad espacial", size=13, space_before=10, space_after=4)
add_para(
    "S(m, n, p) = O(m) por el arreglo de pedidos + O(n·p) por el conteo de "
    "categorías por cliente + O(n) por la lista clientes + O(log n) por la "
    "pila de recursión de QuickSort optimizado en cola.  Resultado: "
    "S(m, n, p) = O(m + n·p).",
    size=11, space_after=8,
)

# Conceptos del curso aplicados
add_heading("5. Algoritmos del curso aplicados", size=13, space_before=4, space_after=4)
add_para(
    "Práctica 4 (Recursividad): funciones recursivas auxiliares en "
    "recursividad.js (suma total, conteo de clientes premium, top cliente) y "
    "la propia estructura recursiva del QuickSort.",
    size=11, space_after=4,
)
add_para(
    "Práctica 5 (Ordenamiento): QuickSort recursivo basado en el esquema "
    "particion_por_Nombre / quickSort_por_Nombre, con mediana de tres como "
    "pivote, insertion sort en particiones pequeñas y recursión optimizada "
    "en cola para garantizar pila O(log n).",
    size=11, space_after=4,
)
add_para(
    "Práctica 6 (Búsqueda): búsqueda binaria recursiva por nombre en "
    "busqueda.js, usada por la utilidad posicionDeCliente(ranking, name) que "
    "consulta en O(log n) la posición de un cliente dentro del ranking.",
    size=11, space_after=8,
)

# Validación
add_heading("6. Verificación con el ejemplo del enunciado", size=13, space_before=4, space_after=4)
add_para(
    'Para el caso "Customer1 Laptop Technology 3000 1 10;Customer2 Shirt '
    'Clothing 50 2 12;Customer1 Mouse Technology 100 1 15;Customer2 Shoes '
    'Clothing 200 1 20;Customer3 TV Technology 2500 1 25", la función '
    "retorna:",
    size=11, space_after=2,
)
ej = doc.add_paragraph()
ej.paragraph_format.left_indent = Cm(0.8)
r = ej.add_run("1) Customer1 3100 Technology\n2) Customer3 2500 Technology\n3) Customer2 300 Clothing")
r.font.name = "Consolas"
r.font.size = Pt(10)

# =============================================================================
# 7. APLICATIVO WEB
# =============================================================================
add_heading("7. Aplicativo web entregado", size=13, space_before=10, space_after=4)
add_para(
    "Además del entregable académico en TareaFinal_EDyA1/, el proyecto "
    "incluye una aplicación web (PWA) desplegada en Vercel y accesible "
    "desde https://tarea-final-e-dy-a1.vercel.app/app/index.html. La app "
    "organiza la solución en cinco paneles: Resultados muestra el ranking "
    "obtenido por calcularPedidos junto con métricas y gráficos; Cargar "
    "caso permite pegar texto, subir un CSV o generar un caso sintético "
    "reproducible con (m, n, p, seed); Benchmark contrasta empíricamente "
    "MergeSort, QuickSort 3-way, RadixSort, Insertion Sort y Array.sort "
    "nativo sobre el mismo dataset; Pruebas ejecuta la suite "
    "automatizada; y Documentación reúne el análisis de complejidad y "
    "las decisiones de diseño.",
    size=11, space_after=6,
)

# =============================================================================
# 8. SUITE DE PRUEBAS
# =============================================================================
add_heading("8. Suite de pruebas automatizada", size=13, space_before=4, space_after=4)
add_para(
    "El archivo app/js/tests.js corre tanto en Node (node app/js/tests.js) "
    "como en la pestaña Pruebas del dashboard, y valida diez escenarios:",
    size=11, space_after=2,
)

tests = [
    "Ejemplo exacto del enunciado.",
    "Empate en totalSpent → desempate por favoriteCategory descendente.",
    "Empate doble → desempate por customer ascendente.",
    "Acumulación correcta de price · quantity.",
    "favoriteCategory equivale a la categoría con más transacciones.",
    "Tolerancia a espacios extra y a ‘;’ final.",
    "Records mal formados se ignoran sin lanzar excepción.",
    "String vacío produce salida vacía.",
    "Volumen alto: 10 000 records procesados sin desbordar la pila.",
    "Estabilidad: orden determinístico ante empates totales.",
]
for t in tests:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(0)
    p.add_run(t).font.size = Pt(10.5)

add_para(
    "Resultado actual: 10/10 pruebas pasan. La integración continua "
    "(.github/workflows/deploy-pages.yml) las ejecuta antes de cada "
    "despliegue, garantizando que sólo versiones verificadas llegan a "
    "producción.",
    size=11, space_after=6,
)

# =============================================================================
# 9. DECISIONES DE DISEÑO
# =============================================================================
add_heading("9. Decisiones de diseño relevantes", size=13, space_before=4, space_after=4)
add_para(
    "Interpretación de favoriteCategory. Se eligió la categoría con mayor "
    "número de transacciones (un record = una compra), no la suma de "
    "quantity ni del gasto. En el ejemplo del enunciado las tres "
    "interpretaciones coinciden, pero ésta es la más simple de defender "
    "y la única que respeta literalmente el texto ‘en la que más "
    "compras realizó’. Empates por frecuencia se rompen escogiendo la "
    "categoría lexicográficamente mayor, consistente con el criterio 2 "
    "del orden global.",
    size=11, space_after=4,
)
add_para(
    "QuickSort frente a MergeSort. El entregable académico usa QuickSort "
    "recursivo siguiendo el esquema particion_por_Nombre de la Práctica "
    "5, con mediana de tres, cambio a Insertion Sort en particiones "
    "pequeñas (umbral 16) y recursión optimizada en cola para mantener "
    "la pila en O(log n). El dashboard usa MergeSort bottom-up estable "
    "para garantizar O(n log n) en el peor caso y permitir comparación "
    "multi-criterio sin desempates secundarios. Ambas variantes resuelven "
    "el mismo enunciado y producen idéntica salida.",
    size=11, space_after=4,
)
add_para(
    "Parser carácter a carácter. Evita el doble split tradicional que "
    "genera del orden de 7·m strings intermedios y presión de "
    "recolección de basura. El scan es O(L) donde L = m·k, con "
    "asignaciones mínimas.",
    size=11, space_after=2,
)

doc.save(OUT)
print(f"OK: {OUT}")
